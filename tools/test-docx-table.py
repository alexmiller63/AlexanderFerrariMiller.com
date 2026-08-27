#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_FILE = ROOT / "test-docx-table.docx"


def set_cell_width(cell, width_twips):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    tc_width = tc_pr.first_child_found_in("w:tcW")

    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)

    tc_width.set(qn("w:w"), str(width_twips))
    tc_width.set(qn("w:type"), "dxa")


def set_table_width(table, width_twips):
    tbl_pr = table._tbl.tblPr

    tbl_width = tbl_pr.first_child_found_in("w:tblW")

    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)

    tbl_width.set(qn("w:w"), str(width_twips))
    tbl_width.set(qn("w:type"), "dxa")


def set_fixed_layout(table):
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.first_child_found_in("w:tblLayout")

    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)

    layout.set(qn("w:type"), "fixed")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr

    borders = tbl_pr.first_child_found_in("w:tblBorders")

    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"

        element = borders.find(qn(tag))

        if element is None:
            element = OxmlElement(tag)
            borders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def main():
    document = Document()

    section = document.sections[0]

    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Letter page:
    # 8.5 inches wide
    # minus 0.75-inch left margin
    # minus 0.75-inch right margin
    # = 7.0 inches usable width
    #
    # Word uses twips:
    # 1 inch = 1440 twips
    page_width_twips = int(section.page_width / 635)
    left_margin_twips = int(section.left_margin / 635)
    right_margin_twips = int(section.right_margin / 635)

    usable_width_twips = (
        page_width_twips
        - left_margin_twips
        - right_margin_twips
    )

    column_width_twips = usable_width_twips // 2

    table = document.add_table(rows=1, cols=2)

    table.autofit = False

    set_fixed_layout(table)
    set_table_width(table, usable_width_twips)
    set_table_borders(table)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    set_cell_width(left_cell, column_width_twips)
    set_cell_width(right_cell, column_width_twips)

    left_cell.width = Inches(3.5)
    right_cell.width = Inches(3.5)

    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    left_paragraph = left_cell.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    left_run = left_paragraph.add_run("Left cell Kilroy was here")
    left_run.font.size = Pt(12)

    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    right_run = right_paragraph.add_run("Right cell Kilroy was here")
    right_run.font.size = Pt(12)

    document.save(OUTPUT_FILE)

    print(f"Created: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()