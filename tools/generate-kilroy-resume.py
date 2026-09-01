#!/usr/bin/env python3
"""Generate period-styled Kilroy resume formats from canonical Markdown."""

from __future__ import annotations

import re
import subprocess
import sys
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)


ROOT = Path(__file__).resolve().parents[1]
RESUME_DIR = ROOT / "kilroy" / "resumes"


def set_font(run, size=8.5, bold=False, italic=False):
    run.font.name = "Courier New"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Courier New")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Courier New")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    p_pr.append(borders)


def add_text(paragraph, text, size=8.5, bold=False, italic=False):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        value = part[2:-2] if is_bold else part
        set_font(paragraph.add_run(value), size=size, bold=bold or is_bold, italic=italic)


def ascii_text(text):
    return text.replace("—", "-").replace("–", "-").replace("·", "-")


def build_docx(source: Path, output: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Courier New"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    normal.font.size = Pt(8.5)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0

    lines = source.read_text(encoding="utf-8").splitlines()
    first_contact_block = True
    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(3)
            add_text(p, line[2:], size=18, bold=True)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(3)
            add_text(p, line[3:].upper(), size=10, bold=True)
            add_bottom_border(p)
            first_contact_block = False
        elif line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(1)
            add_text(p, line[4:].upper(), size=8.8, bold=True)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.28)
            p.paragraph_format.first_line_indent = Inches(-0.14)
            p.paragraph_format.space_after = Pt(1.5)
            add_text(p, line[2:], size=8.2)
        else:
            p = doc.add_paragraph()
            if first_contact_block:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(3)
            add_text(p, line, size=8.5, bold=line.startswith("**") and line.endswith("**"))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("KILROY WAS HERE - TECHNICAL RESUME - 1945"), size=8)
    doc.core_properties.title = "James J. Kilroy - Technical Resume"
    doc.core_properties.author = "Alexander Ferrari Miller"
    doc.save(output)


def build_pdf(source: Path, output: Path):
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "KilroyBody", parent=styles["BodyText"], fontName="Courier",
        fontSize=8.5, leading=10.5, spaceAfter=3,
    )
    title = ParagraphStyle(
        "KilroyTitle", parent=body, fontName="Courier-Bold",
        fontSize=18, leading=20, alignment=TA_CENTER, spaceAfter=4,
    )
    centered = ParagraphStyle(
        "KilroyCentered", parent=body, alignment=TA_CENTER, spaceAfter=4,
    )
    h2 = ParagraphStyle(
        "KilroyH2", parent=body, fontName="Courier-Bold",
        fontSize=10.5, leading=12, spaceBefore=8, spaceAfter=4,
        borderWidth=0, borderPadding=0,
    )
    h3 = ParagraphStyle(
        "KilroyH3", parent=body, fontName="Courier-Bold",
        fontSize=9, leading=10.5, spaceBefore=5, spaceAfter=2,
    )

    story = []
    bullets = []
    before_first_h2 = True

    def flush_bullets():
        nonlocal bullets
        if bullets:
            story.append(ListFlowable(
                [ListItem(Paragraph(item, body)) for item in bullets],
                bulletType="bullet", start="circle", leftIndent=18,
                bulletFontName="Courier", bulletFontSize=7,
                spaceAfter=3,
            ))
            bullets = []

    for raw in source.read_text(encoding="utf-8").splitlines():
        line = ascii_text(raw.strip())
        if not line:
            flush_bullets()
            continue
        markup = escape(line).replace("**", "<b>", 1)
        if "**" in markup:
            markup = markup.replace("**", "</b>", 1)
        if line.startswith("# "):
            story.append(Paragraph(escape(line[2:]), title))
        elif line.startswith("## "):
            flush_bullets()
            story.append(Paragraph(escape(line[3:].upper()), h2))
            before_first_h2 = False
        elif line.startswith("### "):
            flush_bullets()
            story.append(Paragraph(escape(line[4:].upper()), h3))
        elif line.startswith("- "):
            item = escape(line[2:]).replace("**", "<b>", 1)
            if "**" in item:
                item = item.replace("**", "</b>", 1)
            bullets.append(item)
        else:
            flush_bullets()
            story.append(Paragraph(markup, centered if before_first_h2 else body))
    flush_bullets()

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Courier", 7.5)
        canvas.drawCentredString(
            letter[0] / 2, 0.35 * inch,
            f"KILROY WAS HERE - TECHNICAL RESUME - 1945 - PAGE {doc.page}",
        )
        canvas.restoreState()

    document = SimpleDocTemplate(
        str(output), pagesize=letter,
        leftMargin=0.72 * inch, rightMargin=0.72 * inch,
        topMargin=0.50 * inch, bottomMargin=0.48 * inch,
        title="James J. Kilroy - Technical Resume",
        author="Alexander Ferrari Miller",
    )
    document.build(story, onFirstPage=footer, onLaterPages=footer)


def main():
    if len(sys.argv) != 2:
        raise SystemExit("Usage: generate-kilroy-resume.py STEM")
    stem = sys.argv[1]
    source = RESUME_DIR / f"{stem}.md"
    if not source.is_file():
        raise SystemExit(f"Missing source: {source}")

    docx = RESUME_DIR / f"{stem}.docx"
    pdf = RESUME_DIR / f"{stem}.pdf"
    odt = RESUME_DIR / f"{stem}.odt"
    txt = RESUME_DIR / f"{stem}.txt"

    build_docx(source, docx)
    build_pdf(source, pdf)
    subprocess.run(["pandoc", str(source), "-o", str(odt)], check=True)
    subprocess.run(["pandoc", str(source), "-t", "plain", "-o", str(txt)], check=True)
    plain = ascii_text(txt.read_text(encoding="utf-8"))
    txt.write_text(plain.encode("ascii", "replace").decode("ascii"), encoding="ascii")
    if not pdf.is_file():
        raise SystemExit(f"PDF generation failed: {pdf}")

    print(f"Generated {docx.name}, {pdf.name}, {odt.name}, and {txt.name}")


if __name__ == "__main__":
    main()
